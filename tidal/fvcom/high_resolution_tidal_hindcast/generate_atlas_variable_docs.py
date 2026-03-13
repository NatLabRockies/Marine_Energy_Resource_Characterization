"""
Generate Word document documentation for Marine Energy Atlas variables.

This script creates a .docx file documenting the user-facing variables
on the Marine Energy Atlas with:
- Short summary (for info popup on map)
- Complete documentation (full detail)
- Formulas in Word equation format (proper OMML)
- References

User-facing variables:
- Mean Current Speed
- Mean Power Density
- 95th Percentile Current Speed
- Grid Resolution
- Minimum Water Depth
- Maximum Water Depth
"""

from pathlib import Path
from datetime import datetime
import zipfile
import tempfile
import argparse
import re

import numpy as np

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsmap, qn
from docx.oxml import OxmlElement
from lxml import etree

# Import citation manager for BibTeX/CSL formatting
from src.citation_manager import format_reference, format_citation
from src.variable_registry import VARIABLE_REGISTRY

# =============================================================================
# NREL Template Configuration
# =============================================================================

# Path to NREL template (.dotx)
NREL_TEMPLATE_PATH = Path(__file__).parent / "nrel-report-template (1).dotx"

# Style mappings: our generic names -> NREL style names
NREL_STYLES = {
    "body": "NLR_Body_Text",
    "heading1": "NLR_Head_01",
    "heading2": "NLR_Head_02",
    "heading3": "NLR_Head_03",
    "heading4": "NLR_Head_04",
    "bullet": "NLR_Bullet_01",
    "bullet2": "NLR_Bullet_02",
    "equation": "NLR_Equation",
    "reference": "NLR_Reference",
    "title": "NLR_Head_01",  # Use H1 for title
    "caption": "NLR_Figure_Caption",
    "table_caption": "NLR_Table_Caption",
}


def convert_dotx_to_docx(dotx_path):
    """
    Convert a .dotx template to a .docx file that python-docx can open.

    Returns path to temporary .docx file.
    """
    if not Path(dotx_path).exists():
        raise FileNotFoundError(f"Template not found: {dotx_path}")

    # Create temp file for the converted document
    temp_file = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    temp_path = temp_file.name
    temp_file.close()

    # Read original and modify content type
    with zipfile.ZipFile(dotx_path, "r") as zin:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    # Replace template content type with document content type
                    data = data.replace(
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.template.main+xml",
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                    )
                zout.writestr(item, data)

    return temp_path


def get_style(doc, style_key):
    """
    Get the appropriate style from the document.
    Falls back to default styles if NREL styles not available.
    """
    nrel_style = NREL_STYLES.get(style_key)

    # Try NREL style first
    if nrel_style:
        try:
            return doc.styles[nrel_style]
        except KeyError:
            pass

    # Fallback mappings
    fallbacks = {
        "body": "Normal",
        "heading1": "Heading 1",
        "heading2": "Heading 2",
        "heading3": "Heading 3",
        "heading4": "Heading 4",
        "bullet": "List Bullet",
        "bullet2": "List Bullet 2",
        "equation": "Normal",
        "reference": "Normal",
        "title": "Title",
        "caption": "Caption",
    }

    fallback = fallbacks.get(style_key, "Normal")
    try:
        return doc.styles[fallback]
    except KeyError:
        return doc.styles["Normal"]


def add_body_paragraph(doc, text):
    """
    Add a body paragraph with NREL body text style if available.
    Falls back to Normal style if not.
    """
    # Try NREL body text style first
    body_style = NREL_STYLES.get("body", "Normal")
    try:
        return doc.add_paragraph(text, style=body_style)
    except KeyError:
        return doc.add_paragraph(text)


def get_reference_style(doc):
    """
    Get the appropriate reference style from the document.
    Returns NREL reference style if available, otherwise Normal.
    """
    ref_style = NREL_STYLES.get("reference", "Normal")
    try:
        doc.styles[ref_style]
        return ref_style
    except KeyError:
        return "Normal"


# =============================================================================
# OMML Equation Builder
# =============================================================================

# Office Math namespace
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NSMAP = {"m": MATH_NS}


def _m(tag):
    """Create an element in the math namespace."""
    return etree.Element(f"{{{MATH_NS}}}{tag}", nsmap=NSMAP)


def _m_sub(tag):
    """Create a subelement helper."""
    return etree.SubElement


def make_text(text, italic=True, bold=False):
    """Create a math run with text."""
    r = _m("r")

    # Run properties
    if not italic or bold:
        rPr = etree.SubElement(r, f"{{{MATH_NS}}}rPr")
        if not italic:
            sty = etree.SubElement(rPr, f"{{{MATH_NS}}}sty")
            sty.set(f"{{{MATH_NS}}}val", "p")  # plain (not italic)
        if bold:
            sty = etree.SubElement(rPr, f"{{{MATH_NS}}}sty")
            sty.set(f"{{{MATH_NS}}}val", "b")

    t = etree.SubElement(r, f"{{{MATH_NS}}}t")
    t.text = text
    return r


def make_fraction(numerator_elems, denominator_elems):
    """Create a fraction element."""
    f = _m("f")

    # Fraction properties
    fPr = etree.SubElement(f, f"{{{MATH_NS}}}fPr")
    type_elem = etree.SubElement(fPr, f"{{{MATH_NS}}}type")
    type_elem.set(f"{{{MATH_NS}}}val", "bar")  # normal fraction bar

    # Numerator
    num = etree.SubElement(f, f"{{{MATH_NS}}}num")
    for elem in numerator_elems:
        num.append(elem)

    # Denominator
    den = etree.SubElement(f, f"{{{MATH_NS}}}den")
    for elem in denominator_elems:
        den.append(elem)

    return f


def make_subscript(base_elems, sub_elems):
    """Create a subscript element."""
    sSub = _m("sSub")

    # Base
    e = etree.SubElement(sSub, f"{{{MATH_NS}}}e")
    for elem in base_elems:
        e.append(elem)

    # Subscript
    sub = etree.SubElement(sSub, f"{{{MATH_NS}}}sub")
    for elem in sub_elems:
        sub.append(elem)

    return sSub


def make_superscript(base_elems, sup_elems):
    """Create a superscript element."""
    sSup = _m("sSup")

    # Base
    e = etree.SubElement(sSup, f"{{{MATH_NS}}}e")
    for elem in base_elems:
        e.append(elem)

    # Superscript
    sup = etree.SubElement(sSup, f"{{{MATH_NS}}}sup")
    for elem in sup_elems:
        sup.append(elem)

    return sSup


def make_subsup(base_elems, sub_elems, sup_elems):
    """Create a subscript-superscript element."""
    sSubSup = _m("sSubSup")

    # Base
    e = etree.SubElement(sSubSup, f"{{{MATH_NS}}}e")
    for elem in base_elems:
        e.append(elem)

    # Subscript
    sub = etree.SubElement(sSubSup, f"{{{MATH_NS}}}sub")
    for elem in sub_elems:
        sub.append(elem)

    # Superscript
    sup = etree.SubElement(sSubSup, f"{{{MATH_NS}}}sup")
    for elem in sup_elems:
        sup.append(elem)

    return sSubSup


def make_radical(content_elems, degree_elems=None):
    """Create a radical (square root) element."""
    rad = _m("rad")

    # Radical properties
    radPr = etree.SubElement(rad, f"{{{MATH_NS}}}radPr")
    if degree_elems is None:
        # Hide degree for square root
        degHide = etree.SubElement(radPr, f"{{{MATH_NS}}}degHide")
        degHide.set(f"{{{MATH_NS}}}val", "1")

    # Degree (for nth roots)
    deg = etree.SubElement(rad, f"{{{MATH_NS}}}deg")
    if degree_elems:
        for elem in degree_elems:
            deg.append(elem)

    # Content under radical
    e = etree.SubElement(rad, f"{{{MATH_NS}}}e")
    for elem in content_elems:
        e.append(elem)

    return rad


def make_overbar(content_elems):
    """Create an overbar (mean symbol) element."""
    bar = _m("bar")

    # Bar properties
    barPr = etree.SubElement(bar, f"{{{MATH_NS}}}barPr")
    pos = etree.SubElement(barPr, f"{{{MATH_NS}}}pos")
    pos.set(f"{{{MATH_NS}}}val", "top")

    # Content
    e = etree.SubElement(bar, f"{{{MATH_NS}}}e")
    for elem in content_elems:
        e.append(elem)

    return bar


def make_nary(operator, sub_elems=None, sup_elems=None, content_elems=None):
    """Create an n-ary operator (summation, product, etc.)."""
    nary = _m("nary")

    # N-ary properties
    naryPr = etree.SubElement(nary, f"{{{MATH_NS}}}naryPr")

    # Operator character
    chr_elem = etree.SubElement(naryPr, f"{{{MATH_NS}}}chr")
    chr_elem.set(f"{{{MATH_NS}}}val", operator)

    # Limits location (under/over vs subscript/superscript)
    limLoc = etree.SubElement(naryPr, f"{{{MATH_NS}}}limLoc")
    limLoc.set(f"{{{MATH_NS}}}val", "undOvr")

    # Subscript (lower limit)
    sub = etree.SubElement(nary, f"{{{MATH_NS}}}sub")
    if sub_elems:
        for elem in sub_elems:
            sub.append(elem)

    # Superscript (upper limit)
    sup = etree.SubElement(nary, f"{{{MATH_NS}}}sup")
    if sup_elems:
        for elem in sup_elems:
            sup.append(elem)

    # Content being operated on
    e = etree.SubElement(nary, f"{{{MATH_NS}}}e")
    if content_elems:
        for elem in content_elems:
            e.append(elem)

    return nary


def make_delimited(content_elems, left="(", right=")"):
    """Create delimited content (parentheses, brackets, etc.)."""
    d = _m("d")

    # Delimiter properties
    dPr = etree.SubElement(d, f"{{{MATH_NS}}}dPr")

    begChr = etree.SubElement(dPr, f"{{{MATH_NS}}}begChr")
    begChr.set(f"{{{MATH_NS}}}val", left)

    endChr = etree.SubElement(dPr, f"{{{MATH_NS}}}endChr")
    endChr.set(f"{{{MATH_NS}}}val", right)

    # Content
    e = etree.SubElement(d, f"{{{MATH_NS}}}e")
    for elem in content_elems:
        e.append(elem)

    return d


def make_equation_paragraph(elements):
    """Wrap elements in an oMathPara for display."""
    oMathPara = _m("oMathPara")
    oMath = etree.SubElement(oMathPara, f"{{{MATH_NS}}}oMath")
    for elem in elements:
        oMath.append(elem)
    return oMathPara


# =============================================================================
# Equation Definitions for Each Variable
# =============================================================================


def build_mean_current_speed_equation():
    """
    Build: Ū = (1/T) Σ_{t=1}^{T} [(1/N_σ) Σ_{i=1}^{N_σ} U_{i,t}]
    """
    elements = []

    # U with overbar
    elements.append(make_overbar([make_text("U")]))

    # Equals sign
    elements.append(make_text(" = ", italic=False))

    # (1/T)
    elements.append(make_fraction([make_text("1", italic=False)], [make_text("T")]))

    # Space
    elements.append(make_text(" ", italic=False))

    # Outer summation
    outer_sum = make_nary(
        "∑",
        sub_elems=[make_text("t=1", italic=False)],
        sup_elems=[make_text("T")],
        content_elems=[
            make_delimited(
                [
                    # (1/N_σ)
                    make_fraction(
                        [make_text("1", italic=False)],
                        [make_subscript([make_text("N")], [make_text("σ")])],
                    ),
                    make_text(" ", italic=False),
                    # Inner summation
                    make_nary(
                        "∑",
                        sub_elems=[make_text("i=1", italic=False)],
                        sup_elems=[make_subscript([make_text("N")], [make_text("σ")])],
                        content_elems=[
                            make_subscript(
                                [make_text("U")], [make_text("i,t", italic=False)]
                            )
                        ],
                    ),
                ],
                left="[",
                right="]",
            )
        ],
    )
    elements.append(outer_sum)

    return make_equation_paragraph(elements)


def build_mean_power_density_equation():
    """
    Build: P̄ = (1/T) Σ_{t=1}^{T} [(1/N_σ) Σ_{i=1}^{N_σ} (½ρU³_{i,t})]
    """
    elements = []

    # P with overbar
    elements.append(make_overbar([make_text("P")]))

    # Equals sign
    elements.append(make_text(" = ", italic=False))

    # (1/T)
    elements.append(make_fraction([make_text("1", italic=False)], [make_text("T")]))

    elements.append(make_text(" ", italic=False))

    # Outer summation
    outer_sum = make_nary(
        "∑",
        sub_elems=[make_text("t=1", italic=False)],
        sup_elems=[make_text("T")],
        content_elems=[
            make_delimited(
                [
                    make_fraction(
                        [make_text("1", italic=False)],
                        [make_subscript([make_text("N")], [make_text("σ")])],
                    ),
                    make_text(" ", italic=False),
                    make_nary(
                        "∑",
                        sub_elems=[make_text("i=1", italic=False)],
                        sup_elems=[make_subscript([make_text("N")], [make_text("σ")])],
                        content_elems=[
                            make_delimited(
                                [
                                    make_fraction(
                                        [make_text("1", italic=False)],
                                        [make_text("2", italic=False)],
                                    ),
                                    make_text("ρ"),
                                    make_superscript(
                                        [
                                            make_subscript(
                                                [make_text("U")],
                                                [make_text("i,t", italic=False)],
                                            )
                                        ],
                                        [make_text("3", italic=False)],
                                    ),
                                ]
                            )
                        ],
                    ),
                ],
                left="[",
                right="]",
            )
        ],
    )
    elements.append(outer_sum)

    return make_equation_paragraph(elements)


def build_p95_current_speed_equation():
    """
    Build: U_{95} = P_{95}(max_{σ}(U_{i,t}))
    """
    elements = []

    # U_95
    elements.append(make_subscript([make_text("U")], [make_text("95", italic=False)]))

    elements.append(make_text(" = ", italic=False))

    # P_95
    elements.append(make_subscript([make_text("P")], [make_text("95", italic=False)]))

    # (max_σ(U_{i,t}))
    elements.append(
        make_delimited(
            [
                make_subscript([make_text("max")], [make_text("σ")]),
                make_delimited(
                    [make_subscript([make_text("U")], [make_text("i,t", italic=False)])]
                ),
            ]
        )
    )

    return make_equation_paragraph(elements)


def build_grid_resolution_equation():
    """
    Build: R = (1/3)(d₁ + d₂ + d₃)
    """
    elements = []

    elements.append(make_text("R"))
    elements.append(make_text(" = ", italic=False))

    # 1/3
    elements.append(
        make_fraction([make_text("1", italic=False)], [make_text("3", italic=False)])
    )

    # (d₁ + d₂ + d₃)
    elements.append(
        make_delimited(
            [
                make_subscript([make_text("d")], [make_text("1", italic=False)]),
                make_text(" + ", italic=False),
                make_subscript([make_text("d")], [make_text("2", italic=False)]),
                make_text(" + ", italic=False),
                make_subscript([make_text("d")], [make_text("3", italic=False)]),
            ]
        )
    )

    return make_equation_paragraph(elements)


def build_min_depth_equation():
    """
    Build: d_{min} = min_{t}(h + ζ_t)
    """
    elements = []

    # d_min
    elements.append(make_subscript([make_text("d")], [make_text("min", italic=False)]))

    elements.append(make_text(" = ", italic=False))

    # min_t
    elements.append(make_subscript([make_text("min")], [make_text("t")]))

    # (h + ζ_t)
    elements.append(
        make_delimited(
            [
                make_text("h"),
                make_text(" + ", italic=False),
                make_subscript([make_text("ζ")], [make_text("t")]),
            ]
        )
    )

    return make_equation_paragraph(elements)


def build_max_depth_equation():
    """
    Build: d_{max} = max_{t}(h + ζ_t)
    """
    elements = []

    # d_max
    elements.append(make_subscript([make_text("d")], [make_text("max", italic=False)]))

    elements.append(make_text(" = ", italic=False))

    # max_t
    elements.append(make_subscript([make_text("max")], [make_text("t")]))

    # (h + ζ_t)
    elements.append(
        make_delimited(
            [
                make_text("h"),
                make_text(" + ", italic=False),
                make_subscript([make_text("ζ")], [make_text("t")]),
            ]
        )
    )

    return make_equation_paragraph(elements)


def build_speed_equation():
    """
    Build: U_{i,t} = √(u² + v²)
    """
    elements = []

    # U_{i,t}
    elements.append(make_subscript([make_text("U")], [make_text("i,t", italic=False)]))

    elements.append(make_text(" = ", italic=False))

    # √(u² + v²)
    elements.append(
        make_radical(
            [
                make_superscript([make_text("u")], [make_text("2", italic=False)]),
                make_text(" + ", italic=False),
                make_superscript([make_text("v")], [make_text("2", italic=False)]),
            ]
        )
    )

    return make_equation_paragraph(elements)


# Map variable keys to their equation builders
EQUATION_BUILDERS = {
    "mean_current_speed": build_mean_current_speed_equation,
    "mean_power_density": build_mean_power_density_equation,
    "p95_current_speed": build_p95_current_speed_equation,
    "grid_resolution": build_grid_resolution_equation,
    "min_water_depth": build_min_depth_equation,
    "max_water_depth": build_max_depth_equation,
}


# =============================================================================
# "Where" Variable Definition Builders
# =============================================================================


def build_velocity_magnitude_definition():
    """Build: U_{i,t} = √(u_{i,t}² + v_{i,t}²)"""
    elements = []
    elements.append(make_subscript([make_text("U")], [make_text("i,t", italic=False)]))
    elements.append(make_text(" = ", italic=False))
    elements.append(
        make_radical(
            [
                make_superscript(
                    [
                        make_subscript(
                            [make_text("u")], [make_text("i,t", italic=False)]
                        )
                    ],
                    [make_text("2", italic=False)],
                ),
                make_text(" + ", italic=False),
                make_superscript(
                    [
                        make_subscript(
                            [make_text("v")], [make_text("i,t", italic=False)]
                        )
                    ],
                    [make_text("2", italic=False)],
                ),
            ]
        )
    )
    return make_equation_paragraph(elements)


def build_power_density_definition():
    """Build: P_{i,t} = ½ρU³_{i,t}"""
    elements = []
    elements.append(make_subscript([make_text("P")], [make_text("i,t", italic=False)]))
    elements.append(make_text(" = ", italic=False))
    elements.append(
        make_fraction([make_text("1", italic=False)], [make_text("2", italic=False)])
    )
    elements.append(make_text("ρ"))
    elements.append(
        make_superscript(
            [make_subscript([make_text("U")], [make_text("i,t", italic=False)])],
            [make_text("3", italic=False)],
        )
    )
    return make_equation_paragraph(elements)


def build_n_sigma_definition():
    """Build: N_σ = 10"""
    elements = []
    elements.append(make_subscript([make_text("N")], [make_text("σ")]))
    elements.append(make_text(" = 10", italic=False))
    return make_equation_paragraph(elements)


def build_rho_definition():
    """Build: ρ = 1025 kg/m³"""
    elements = []
    elements.append(make_text("ρ"))
    elements.append(make_text(" = 1025 kg/m³", italic=False))
    return make_equation_paragraph(elements)


def build_T_definition():
    """Build: T = 1 year"""
    elements = []
    elements.append(make_text("T"))
    elements.append(make_text(" = 1 year", italic=False))
    return make_equation_paragraph(elements)


def build_h_definition():
    """Build: h = bathymetry depth"""
    elements = []
    elements.append(make_text("h"))
    elements.append(make_text(" = bathymetry depth below NAVD88 (m)", italic=False))
    return make_equation_paragraph(elements)


def build_zeta_definition():
    """Build: ζ_t = sea surface elevation"""
    elements = []
    elements.append(make_subscript([make_text("ζ")], [make_text("t")]))
    elements.append(
        make_text(" = sea surface elevation above NAVD88 at time t (m)", italic=False)
    )
    return make_equation_paragraph(elements)


def build_d_edges_definition():
    """Build: d₁, d₂, d₃ = edge lengths"""
    elements = []
    elements.append(make_subscript([make_text("d")], [make_text("1", italic=False)]))
    elements.append(make_text(", ", italic=False))
    elements.append(make_subscript([make_text("d")], [make_text("2", italic=False)]))
    elements.append(make_text(", ", italic=False))
    elements.append(make_subscript([make_text("d")], [make_text("3", italic=False)]))
    elements.append(
        make_text(" = geodesic distances between triangle vertices (m)", italic=False)
    )
    return make_equation_paragraph(elements)


def build_max_sigma_definition():
    """Build: max_σ = maximum across sigma layers"""
    elements = []
    elements.append(make_subscript([make_text("max")], [make_text("σ")]))
    elements.append(
        make_text(
            " = maximum value across all sigma layers at each timestep", italic=False
        )
    )
    return make_equation_paragraph(elements)


def build_p95_definition():
    """Build: P₉₅ = 95th percentile"""
    elements = []
    elements.append(make_subscript([make_text("P")], [make_text("95", italic=False)]))
    elements.append(
        make_text(" = 95th percentile operator over the time series", italic=False)
    )
    return make_equation_paragraph(elements)


# Define which variable definitions to show for each variable
VARIABLE_DEFINITIONS = {
    "mean_current_speed": [
        ("velocity_magnitude", None),
        ("u_component", None),
        ("v_component", None),
        ("n_sigma", None),
        ("T", None),
    ],
    "mean_power_density": [
        ("power_density", None),
        ("rho", None),
        ("velocity_magnitude", None),
        ("u_component", None),
        ("v_component", None),
        ("n_sigma", None),
        ("T", None),
    ],
    "p95_current_speed": [
        ("velocity_magnitude", None),
        ("u_component", None),
        ("v_component", None),
        ("max_sigma", None),
        ("p95", None),
        ("n_sigma", None),
        ("T", None),
    ],
    "grid_resolution": [
        ("d_edges", None),
    ],
    "min_water_depth": [
        ("h", None),
        ("zeta", None),
        ("T", "hindcast duration"),
    ],
    "max_water_depth": [
        ("h", None),
        ("zeta", None),
        ("T", "hindcast duration"),
    ],
}

DEFINITION_BUILDERS = {
    "velocity_magnitude": build_velocity_magnitude_definition,
    "power_density": build_power_density_definition,
    "n_sigma": build_n_sigma_definition,
    "rho": build_rho_definition,
    "T": build_T_definition,
    "h": build_h_definition,
    "zeta": build_zeta_definition,
    "d_edges": build_d_edges_definition,
    "max_sigma": build_max_sigma_definition,
    "p95": build_p95_definition,
}


# =============================================================================
# Variable Specifications
# =============================================================================

ATLAS_VARIABLES = {
    k: v for k, v in VARIABLE_REGISTRY.items() if v["included_on_atlas"]
}

# =============================================================================
# References - BibTeX keys from references/references.bib
# Citations are formatted using citation_manager.py with Chicago CSL
# =============================================================================

# BibTeX keys used in this document (from references/references.bib)
# Keys must match exactly (case-sensitive) with references/references.bib
BIBTEX_KEYS = [
    "iec_62600_201",
    "fvcom",
    "mhkdr_submission",
    "ak_aleutian_spicer2025_spatially",
    "ak_cook_deb2025_characterizing",
    "ak_southeast_brand_2025_tidal",
    "me_western_deb2023_turbulence",
    "me_western_yang2020_modeling",
    "nh_piscataqua_spicer2023_tidal",
    "wa_puget_deb2024_tidal_iec",
    "wa_puget_spicer2024_localized",
    "wa_puget_yang2021_tidal",
]

# Location-specific validation references (using BibTeX keys)
# Keys must match exactly (case-sensitive) with references/references.bib
LOCATION_REFERENCES = {
    "Alaska, Aleutian Islands": ["ak_aleutian_spicer2025_spatially"],
    "Alaska, Cook Inlet": ["ak_cook_deb2025_characterizing"],
    # "Alaska, Southeast": ["ak_southeast_brand_2025_tidal"],
    "Maine, Western Passage": [
        "me_western_deb2023_turbulence",
        "me_western_yang2020_modeling",
    ],
    "New Hampshire, Piscataqua River": ["nh_piscataqua_spicer2023_tidal"],
    "Washington, Puget Sound": [
        "wa_puget_deb2024_tidal_iec",
        "wa_puget_spicer2024_localized",
        "wa_puget_yang2021_tidal",
    ],
}


# =============================================================================
# Document Generation Functions
# =============================================================================


def add_heading(doc, text, level=1, use_nrel_styles=True):
    """
    Add a heading with appropriate formatting.

    If use_nrel_styles is True and NREL styles are available, uses NLR_Head_0X styles.
    Otherwise falls back to built-in Heading X styles.
    """
    if use_nrel_styles:
        style_key = f"heading{level}"
        nrel_style = NREL_STYLES.get(style_key)
        if nrel_style:
            try:
                # Use NREL style by adding paragraph with that style
                p = doc.add_paragraph(text, style=nrel_style)
                return p
            except KeyError:
                pass

    # Fallback to built-in heading styles
    heading = doc.add_heading(text, level=level)
    return heading


def add_omml_equation(doc, omml_element):
    """Add an OMML equation to the document."""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Convert lxml element to OxmlElement for python-docx compatibility
    omml_str = etree.tostring(omml_element, encoding="unicode")
    omml_docx = etree.fromstring(omml_str)

    paragraph._p.append(omml_docx)
    return paragraph


def add_where_section(doc, var_key, compact=False):
    """Add a 'Where:' section with variable definitions as bullet list."""
    if var_key not in VARIABLE_DEFINITIONS:
        return

    definitions = VARIABLE_DEFINITIONS[var_key]

    # Bold "Where:" header, left aligned
    p = doc.add_paragraph()
    p.add_run("Where:").bold = True

    # Text definitions for bullet list (using Unicode for symbols)
    TEXT_DEFINITIONS = {
        "velocity_magnitude": "U = √(u² + v²) — velocity magnitude at sigma layer i at time t (m/s)",
        "u_component": "u = eastward velocity component (m/s), positive toward true east",
        "v_component": "v = northward velocity component (m/s), positive toward true north",
        "power_density": "P = ½ρU³ — power density at sigma layer i at time t (W/m²)",
        "n_sigma": "Nσ = 10 sigma layers (terrain-following vertical layers)",
        "rho": "ρ = 1025 kg/m³ (nominal seawater density)",
        "T": "T = 1 year of hindcast data",
        "h": "h = bathymetry depth below NAVD88 (m)",
        "zeta": "ζₜ = sea surface elevation above NAVD88 at time t (m)",
        "d_edges": "d₁, d₂, d₃ = geodesic distances between triangle vertices (m)",
        "max_sigma": "maxσ = maximum value across all sigma layers at each timestep",
        "p95": "P₉₅ = 95th percentile operator over the time series",
    }

    # Add each definition as a bullet point (left-aligned)
    # Try NREL bullet style first, fall back to List Bullet
    bullet_style = NREL_STYLES.get("bullet", "List Bullet")
    try:
        doc.styles[bullet_style]
    except KeyError:
        bullet_style = "List Bullet"

    for def_key, description_override in definitions:
        if def_key in TEXT_DEFINITIONS:
            doc.add_paragraph(TEXT_DEFINITIONS[def_key], style=bullet_style)


def create_variable_short_section(doc, var_key, var_data):
    """Create the short summary section for a variable (H2 under variable H1).

    This represents the info popup content on the Atlas - just a 1-liner with link.
    """
    # H2: Short Description
    add_heading(doc, "Short Description (i) on Atlas", level=2)

    # 1-line description (same as executive summary table)
    one_liner = var_data.get("one_liner", "")
    p = doc.add_paragraph()
    p.add_run(f"{one_liner} [{var_data['units']}]")

    # Link to detailed documentation (online)
    doc.add_paragraph()
    p = doc.add_paragraph()
    doc_url = var_data.get("documentation_url", "")
    if doc_url:
        p.add_run("For complete methodology, formulas, and references, see: ")
        add_hyperlink(p, var_data["display_name"], doc_url)
    else:
        p.add_run(
            f"See Detailed Documentation below for complete {var_data['display_name']} "
            "methodology, formulas, and references."
        ).italic = True

    doc.add_paragraph()  # Spacing


# DOI lookup for creating hyperlinks (must match BibTeX keys)
REFERENCE_DOIS = {
    "ak_aleutian_spicer2025_spatially": "https://doi.org/10.1016/j.renene.2025.123564",
    "ak_cook_deb2025_characterizing": "https://doi.org/10.1016/j.renene.2025.123345",
    "ak_southeast_brand_2025_tidal": "https://doi.org/10.1016/j.renene.2025.122617",
    "me_western_deb2023_turbulence": "https://doi.org/10.1016/j.renene.2023.04.100",
    "me_western_yang2020_modeling": "https://doi.org/10.3390/jmse8060411",
    "nh_piscataqua_spicer2023_tidal": "https://doi.org/10.3389/fmars.2023.1268348",
    "wa_puget_deb2024_tidal_iec": "https://doi.org/10.1016/j.renene.2024.120767",
    "wa_puget_yang2021_tidal": "https://doi.org/10.1016/j.renene.2021.03.028",
    "wa_puget_spicer2024_localized": "https://doi.org/10.1029/2023JC020401",
    "fvcom": "https://doi.org/10.5670/oceanog.2006.92",
    "iec_62600_201": "https://webstore.iec.ch/en/publication/22099",
    "mhkdr_submission": "https://mhkdr.openei.org/submissions/632",
}


def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to a paragraph.

    python-docx doesn't have native hyperlink support, so we use oxml.
    """
    import re as re_module

    # Get the document part
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    # Create the hyperlink element
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", r_id
    )

    # Create a run for the hyperlink text
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Add hyperlink styling (blue, underlined)
    color = OxmlElement("w:color")
    color.set(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "0563C1"
    )
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "single")
    rPr.append(u)

    new_run.append(rPr)

    # Add the text
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def add_reference_with_doi_link(paragraph, reference_text):
    """
    Add reference text to a paragraph, making any DOI URL a clickable hyperlink.

    Parses the reference text for DOI URLs (https://doi.org/...) and splits
    the text so the DOI becomes a clickable link.

    Handles doubled DOI URLs like "https://doi.org/https://doi.org/10.xxx" that
    occur when BibTeX doi field contains the full URL instead of just the DOI.
    """
    import re as re_module

    # First, fix doubled DOI URLs in the text
    # Pattern: https://doi.org/https://doi.org/10.xxx -> https://doi.org/10.xxx
    fixed_text = re_module.sub(
        r"https?://doi\.org/https?://doi\.org/", "https://doi.org/", reference_text
    )

    # Pattern to match DOI URLs
    # Captures the DOI URL, stopping at whitespace or end of string
    doi_pattern = r"(https?://doi\.org/[^\s]+?)(?:\.?\s*$|\s)"

    match = re_module.search(doi_pattern, fixed_text)

    if match:
        doi_url = match.group(1)
        # Clean up any trailing punctuation from the DOI
        doi_url = doi_url.rstrip(".,;:")

        # Find where the DOI starts in the fixed text
        doi_start = fixed_text.find(doi_url)

        if doi_start > 0:
            # Text before the DOI
            before_text = fixed_text[:doi_start]
            paragraph.add_run(before_text)

        # Add the DOI as a hyperlink
        add_hyperlink(paragraph, doi_url, doi_url)

        # Text after the DOI (if any, excluding trailing punctuation)
        after_start = doi_start + len(doi_url)
        if after_start < len(fixed_text):
            after_text = fixed_text[after_start:].lstrip(".,;: ")
            if after_text:
                paragraph.add_run(" " + after_text)
    else:
        # No DOI found, just add the text normally
        paragraph.add_run(fixed_text)


# =============================================================================
# Table Styling Functions
# =============================================================================

# NREL brand colors and fonts for table styling
NREL_HEADER_BG = "0079C2"  # NREL Blue
NREL_STRIPE_BG = "F0F7FC"  # Light blue for striping
NREL_BORDER_COLOR = "5C6670"  # NREL Gray
NREL_HEADER_FONT = "Arial"  # NREL header font (matches NLR heading styles)


def set_cell_shading(cell, color_hex):
    """Set the background shading color for a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_table_borders(table, color_hex=NREL_BORDER_COLOR, size="4"):
    """
    Set borders on all cells in a table.

    Args:
        table: The python-docx Table object
        color_hex: Border color in hex (without #)
        size: Border width in eighths of a point (4 = 0.5pt, 8 = 1pt)
    """
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")

    # Create table borders element
    tblBorders = OxmlElement("w:tblBorders")

    # Define border sides
    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:color"), color_hex)
        border.set(qn("w:space"), "0")
        tblBorders.append(border)

    tblPr.append(tblBorders)

    # Ensure tblPr is in the table
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)


def add_table_caption(doc, caption_text, table_number=None):
    """
    Add a table caption with NREL styling.

    Args:
        doc: The Document object
        caption_text: The caption text
        table_number: Optional table number (auto-numbered if None)

    Returns:
        The caption paragraph
    """
    # Try NREL table caption style first, then figure caption, then Caption
    caption_style = NREL_STYLES.get("table_caption")
    try:
        p = doc.add_paragraph(style=caption_style)
    except KeyError:
        # Try figure caption as fallback
        try:
            p = doc.add_paragraph(style=NREL_STYLES.get("caption", "Caption"))
        except KeyError:
            p = doc.add_paragraph()

    # Add "Table X. " prefix in bold if table number provided
    if table_number is not None:
        p.add_run(f"Table {table_number}. ").bold = True
    else:
        p.add_run("Table 1. ").bold = True

    # Add caption text
    p.add_run(caption_text)

    return p


def add_figure_caption(doc, caption_text, figure_number=None):
    """
    Add a figure caption with NREL styling.

    Args:
        doc: The Document object
        caption_text: The caption text
        figure_number: Optional figure number (auto-numbered if None)

    Returns:
        The caption paragraph
    """
    caption_style = NREL_STYLES.get("caption", "Caption")
    try:
        p = doc.add_paragraph(style=caption_style)
    except KeyError:
        p = doc.add_paragraph()

    # Add "Figure X. " prefix in bold
    if figure_number is not None:
        p.add_run(f"Figure {figure_number}. ").bold = True
    else:
        p.add_run("Figure 1. ").bold = True

    # Add caption text
    p.add_run(caption_text)

    return p


def render_image_from_spec(doc, image_spec, figure_number=None):
    """
    Render an image from a specification dictionary.

    This function adds an image to the document with a caption below it.
    Use this for consistent image/figure generation throughout the document.

    Args:
        doc: The Document object
        image_spec: Dictionary containing image specification with keys:
            - "path": str - Path to the image file
            - "caption": str - Figure caption text
            - "width": Inches (optional) - Image width (default: 6 inches)
        figure_number: Optional figure number for caption (auto-numbered if None)

    Returns:
        The image paragraph

    Example image_spec:
        {
            "path": "output/sigma_timeseries.png",
            "caption": "Sigma layer depths varying over a tidal cycle",
            "width": Inches(6.5),
        }
    """
    from pathlib import Path

    image_path = image_spec["path"]
    caption = image_spec.get("caption", "")
    width = image_spec.get("width", Inches(6))

    # Check if image exists
    if not Path(image_path).exists():
        # Add placeholder text if image doesn't exist
        p = doc.add_paragraph()
        p.add_run(f"[Image not found: {image_path}]").italic = True
        if caption:
            add_figure_caption(doc, caption, figure_number=figure_number)
        return p

    # Add the image centered
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=width)

    # Add caption below
    if caption:
        add_figure_caption(doc, caption, figure_number=figure_number)

    return p


def create_styled_table(doc, rows, cols, header_row=True):
    """
    Create a table with NREL styling (borders and optional row striping).

    Args:
        doc: The Document object
        rows: Number of rows (including header)
        cols: Number of columns
        header_row: If True, style first row as header with dark background

    Returns:
        The Table object
    """
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set table borders
    set_table_borders(table)

    return table


def style_table_with_striping(table, header_row=True):
    """
    Apply NREL-style formatting to a table with header and row striping.

    Args:
        table: The python-docx Table object
        header_row: If True, style first row as header
    """
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            if row_idx == 0 and header_row:
                # Header row: dark background, white/bold text, header font
                set_cell_shading(cell, NREL_HEADER_BG)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(255, 255, 255)
                        run.font.name = NREL_HEADER_FONT
            elif row_idx % 2 == 0 and header_row:
                # Even data rows (row 2, 4, 6... which are indices 1, 3, 5...
                # after header): no shading (white)
                pass
            elif row_idx % 2 == 1 and header_row:
                # Odd data rows: light stripe
                set_cell_shading(cell, NREL_STRIPE_BG)


def set_cell_text(cell, text, font_name="Arial", font_size=Pt(10), bold=False):
    """Set cell text with consistent font styling."""
    para = cell.paragraphs[0]
    para.clear()
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = font_size - 1
    run.bold = bold
    return run


def render_table_from_spec(doc, table_spec, table_number=None):
    """
    Render a table from a specification dictionary.

    This function creates a styled table with NREL formatting from a declarative
    table specification. Use this for consistent table generation throughout
    the document.

    Args:
        doc: The Document object
        table_spec: Dictionary containing table specification with keys:
            - "caption": str - Table caption text
            - "headers": list[str] - Column header labels
            - "data": dict[str, list] - Column-oriented data where keys are
              column names (matching headers) and values are lists of cell values.
              This format is preferred for easier data entry and maintenance.
            - "column_widths": list[Inches] (optional) - Column widths
        table_number: Optional table number for caption (auto-numbered if None)

    Returns:
        The Table object

    Example table_spec (column-oriented data format):
        {
            "caption": "Sigma Levels and Layer Centers",
            "headers": ["Level Index", "σ Value", "Description"],
            "data": {
                "Level Index": ["0", "1", "2", ...],
                "σ Value": ["0.0", "-0.1", "-0.2", ...],
                "Description": ["Sea surface", "", "", ...],
            },
            "column_widths": [Inches(1.0), Inches(1.0), Inches(3.0)]
        }
    """
    headers = table_spec["headers"]
    data = table_spec["data"]
    caption = table_spec.get("caption", "")
    column_widths = table_spec.get("column_widths")

    # Convert column-oriented data (dict) to row-oriented data (list of lists)
    if isinstance(data, dict):
        # Get number of rows from first column
        first_col = data[headers[0]]
        num_data_rows = len(first_col)
        # Transpose: create rows from columns
        row_data = []
        for row_idx in range(num_data_rows):
            row = [str(data[header][row_idx]) for header in headers]
            row_data.append(row)
        data = row_data

    num_cols = len(headers)
    num_rows = len(data) + 1  # +1 for header row

    # Create the table
    table = create_styled_table(doc, rows=num_rows, cols=num_cols, header_row=True)

    # Populate header row
    header_cells = table.rows[0].cells
    for idx, header_text in enumerate(headers):
        header_cells[idx].text = header_text

    # Set column widths if provided
    if column_widths:
        for idx, width in enumerate(column_widths):
            header_cells[idx].width = width

    # Populate data rows
    for row_idx, row_data in enumerate(data, start=1):
        row = table.rows[row_idx]
        for col_idx, cell_text in enumerate(row_data):
            set_cell_text(row.cells[col_idx], str(cell_text))

    # Apply column widths to all rows
    if column_widths:
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                cell.width = column_widths[idx]

    # Apply NREL styling with header and row striping
    style_table_with_striping(table, header_row=True)

    # Add table caption below the table
    if caption:
        add_table_caption(doc, caption, table_number=table_number)

    return table


# =============================================================================
# Definitions Section
# =============================================================================

# Sigma levels (boundaries): 0 at surface, -1 at seafloor (FVCOM convention)
N_SIGMA_LEVELS = 11
N_SIGMA_LAYERS = 10
_sigma_levels = np.linspace(0, -1.0, N_SIGMA_LEVELS)
_sigma_layers = np.linspace(-0.05, -0.95, N_SIGMA_LAYERS)

sigma_levels_string = str([f"{float(x):.1f}" for x in _sigma_levels]).replace("'", "")
sigma_layers_string = str([f"{float(x):.2f}" for x in _sigma_layers]).replace("'", "")

# Table data for sigma levels (generated from spec)
_sigma_level_indices = [str(i) for i in range(N_SIGMA_LEVELS)]
_sigma_level_values = [
    f"{x:.1f}" if x >= 0 else f"−{abs(x):.1f}" for x in _sigma_levels
]
_sigma_level_positions = (
    ["Sea surface"] + [""] * 4 + ["Mid-depth"] + [""] * 4 + ["Seafloor"]
)

# Table data for sigma layers (generated from spec)
_sigma_layer_indices = [str(i + 1) for i in range(N_SIGMA_LAYERS)]
_sigma_layer_values = [f"−{abs(x):.2f}" for x in _sigma_layers]
_sigma_layer_descriptions = (
    ["Near-surface layer"]
    + [""] * 4
    + ["Mid-depth layer"]
    + [""] * 3
    + ["Near-bottom layer"]
)

# Real-world example: Cook Inlet location (face=125262)
# Data from: AK_cook_inlet.tidal_hindcast_fvcom.face=125262.lat=60.7416496.lon=-151.4299316
_example_h = 26.73  # Bathymetry (m below NAVD88)
_example_zeta_low = -4.95  # Sea surface elevation at low tide (m)
_example_zeta_high = 3.82  # Sea surface elevation at high tide (m)
_example_D_low = _example_h + _example_zeta_low  # Total water depth at low tide
_example_D_high = _example_h + _example_zeta_high  # Total water depth at high tide

# Calculate sigma layer depths at low and high tide
_example_depths_low = [f"{-_example_D_low * sigma:.1f}" for sigma in _sigma_layers]
_example_depths_high = [f"{-_example_D_high * sigma:.1f}" for sigma in _sigma_layers]
_example_depth_diff = [
    f"+{(-_example_D_high * sigma) - (-_example_D_low * sigma):.1f}"
    for sigma in _sigma_layers
]

# Technical terms and their definitions
TECHNICAL_DEFINITIONS = {
    "Hindcast": (
        "A historical simulation of ocean conditions using a numerical model driven by "
        "observed atmospheric forcing and tidal boundary conditions. Unlike a forecast "
        "(which predicts future conditions) or direct measurements, a hindcast reconstructs "
        "past conditions by running the model over a historical time period. The results "
        "represent modeled estimates of what conditions were, not direct observations."
    ),
    "Sigma Layer": {
        "text": [
            (
                "A terrain-following vertical coordinate system where the water column is divided "
                "into layers that conform to both the seafloor (bottom) and the time-varying sea "
                "surface (top). Sigma (σ) expresses vertical position as a proportion of total "
                "water depth, ranging from σ = 0 at the surface to σ = −1 at the seafloor. "
                "Layers stretch and compress dynamically as water depth varies with tides."
            ),
            (
                "Configuration: The FVCOM model uses 10 uniformly spaced sigma layers. Model "
                "variables (velocity, etc.) are computed at each layer center. The physical depth "
                "of any sigma value is: depth = −D × σ, where D = h + ζ (bathymetry + sea surface "
                "elevation). The visualization below illustrates how sigma layers vary with tides."
            ),
        ],
        "images": [
            {
                "path": "output/sigma_timeseries.png",
                "caption": (
                    "Sigma coordinate system over a 4-day tidal cycle at Cook Inlet, Alaska. "
                    "Top panel shows sigma level boundaries (solid) and layer centers (dashed) "
                    "in elevation relative to NAVD88. The seafloor remains fixed while the sea "
                    "surface oscillates with tides, causing sigma layers to expand during high "
                    "tide and compress during low tide."
                ),
                "width": Inches(6.5),
            },
        ],
    },
    "Depth-Averaged": (
        "A value computed by averaging a quantity across all vertical layers (sigma layers) "
        "at a given horizontal location and time. Depth-averaging produces a single "
        "representative value for the entire water column, useful for characterizing overall "
        "flow conditions while accounting for vertical velocity profiles."
    ),
    "Depth-Maximum": (
        "The maximum value of a quantity across all vertical layers (sigma layers) at a given "
        "horizontal location and time. The depth-maximum captures the peak value occurring "
        "anywhere in the water column, which is relevant for structural loading calculations "
        "where components must withstand the highest forces at any depth."
    ),
    "Free-Stream Velocity": (
        "The undisturbed flow velocity that would exist in the absence of any energy "
        "extraction device. All velocity and power density values in this dataset represent "
        "free-stream (undisturbed) conditions and should not be used directly for turbine "
        "array yield estimation. Actual flow through turbine arrays will be modified by: "
        "(1) blockage effects that can reduce channel flow by 10-30% depending on the "
        "blockage ratio; (2) wake interactions where downstream turbines experience 20-40% "
        "velocity deficits in the near-wake region; and (3) device-induced turbulence that "
        "affects fatigue loading. Array yield calculations require site-specific wake "
        "modeling and cannot be derived directly from free-stream resource data."
    ),
    "Unstructured Grid": {
        "text": [
            (
                "A computational mesh composed of irregularly-shaped elements (triangles in FVCOM) "
                "that can vary in size across the domain. This allows higher resolution in areas of "
                "interest (narrow channels, complex coastlines) and coarser resolution in open water, "
                "providing computational efficiency while maintaining accuracy where needed. For the duration of the model run, the horizontal coordinates (lat/lon) of grid remains fixed in space."
            ),
            (
                "Grid Structure: The FVCOM model uses an unstructured triangular mesh where each "
                "element (cell) is defined by three nodes. Model variables are computed at element "
                "centers, while the mesh geometry is defined by node positions. The flexibility of "
                "triangular elements allows the mesh to conform to complex coastlines and bathymetric "
                "features without requiring uniform grid spacing."
            ),
            (
                "Spatial Resolution: Grid resolution varies across the domain based on local "
                "requirements. In narrow channels and near coastlines where currents are strongest "
                "and bathymetry changes rapidly, triangles are smaller (higher resolution). In open "
                "water regions where conditions vary more gradually, triangles are larger (coarser "
                "resolution), reducing computational cost without sacrificing accuracy in areas of "
                "interest. The grid resolution variable in the dataset reports the average edge "
                "length of each triangular element."
            ),
        ],
        "images": [
            {
                "path": "output/unstructured_grid_overview.png",
                "caption": (
                    "Wide angle and detailed view of the underlying FVCOM unstructured triangular mesh used for the High Resolution Tidal Hindcast in Cook Inlet, Alaska. "
                    "The main panel shows the full model domain with 392,002 triangular elements "
                    "conforming to the complex coastline geometry. The inset shows a "
                    "detailed view of the mesh structure, illustrating how element size varies with "
                    "higher resolution near coastlines and in channels where accurate representation "
                    "of tidal dynamics requires finer spatial detail."
                ),
                "width": Inches(6.5),
            },
        ],
    },
    "Model Limitations": {
        "text": [
            (
                "This dataset is derived from numerical model simulations and has inherent "
                "limitations that users should understand when applying the data."
            ),
            (
                "Physics Not Included: The model does not include wave-current interaction "
                "(waves are small in the study domains), atmospheric forcing (wind and pressure "
                "effects on tidal currents are negligible), density-driven estuarine flow "
                "(temperature and salinity effects are small), or storm surge (only astronomical "
                "tidal forcing from 12 constituents via OSU TPXO Tide Models is applied)."
            ),
            (
                "Temporal Limitations: Results represent a single hindcast year and do not capture "
                "interannual variability. Tidal patterns are highly predictable, but extreme values "
                "may vary between years due to meteorological effects."
            ),
            (
                "Spatial Limitations: The model uses wetting/drying algorithms for intertidal zones; "
                "results in these areas should be interpreted with care. Model accuracy may be "
                "reduced near open ocean boundaries. Features smaller than approximately 2-3 times "
                "the local grid spacing cannot be accurately resolved."
            ),
            (
                "Known Data Gaps: Puget Sound is missing December 31, 2015. Specific files were "
                "excluded due to quality issues (see full dataset documentation for details)."
            ),
        ],
    },
}

# Mathematical symbols and their definitions
SYMBOL_DEFINITIONS = {
    "U": "Current speed (velocity magnitude), m/s",
    "Ū": "Mean (time-averaged) current speed, m/s",
    "u": "Eastward velocity component, m/s (positive toward true east)",
    "v": "Northward velocity component, m/s (positive toward true north)",
    "P": "Power density (kinetic energy flux per unit area), W/m²",
    "P̄": "Mean (time-averaged) power density, W/m²",
    "ρ": "Seawater density, kg/m³ (nominal value: 1025 kg/m³)",
    "h": (
        "Bathymetry depth below NAVD88, m (positive downward). A value of h = 20 m means "
        "the seafloor is 20 m below the NAVD88 datum. Negative values (h < 0) indicate "
        "the seafloor is above NAVD88."
    ),
    "ζ": (
        "Sea surface elevation relative to NAVD88, m (positive upward). Following FVCOM "
        "convention, this document uses ζ (zeta); some oceanographic literature uses η (eta) "
        "for the same quantity. Total water depth is calculated as D = h + ζ."
    ),
    "d": "Water depth (total water column height), m",
    "R": "Grid resolution (average triangle edge length), m",
    "T": "Time period (hindcast duration = 1 year)",
    "t": "Time index",
    "i": "Sigma layer index (1 to Nσ)",
    "Nσ": "Number of sigma layers (= 10 in this dataset)",
    "σ": "Sigma coordinate (terrain-following vertical coordinate)",
    "P₉₅": "95th percentile operator",
}

# Acronyms and their definitions
ACRONYM_DEFINITIONS = {
    "FVCOM": (
        "Finite Volume Community Ocean Model. A three-dimensional, unstructured-grid, "
        "finite-volume coastal ocean model developed by the University of Massachusetts "
        "Dartmouth and Woods Hole Oceanographic Institution. FVCOM solves the momentum, "
        "continuity, temperature, salinity, and density equations using terrain-following "
        "sigma coordinates."
    ),
    "NAVD88": (
        "North American Vertical Datum of 1988. A fixed vertical reference datum used for "
        "measuring elevations in North America. Sea surface elevations and bathymetry in "
        "the original model outputs are referenced to NAVD88. Note that NAVD88 differs from "
        "Mean Sea Level (MSL) by a location-dependent offset."
    ),
    "MSL": (
        "Mean Sea Level. In this dataset, MSL is defined as the temporal mean sea surface "
        "height computed from the 1-year hindcast period. Surface elevation values are "
        "reported relative to this MSL so that they oscillate around zero, which is the "
        "standard oceanographic convention expected by most analysis methods. This "
        "dataset-derived MSL may differ from official NOAA tidal datums, which are computed "
        "over longer (19-year) tidal epochs."
    ),
    "IEC": (
        "International Electrotechnical Commission. An international standards organization "
        "that develops and publishes standards for electrical and electronic technologies, "
        "including marine energy systems."
    ),
    "IEC 62600-201": (
        "International standard for tidal energy resource assessment and characterization. "
        "Defines Stage 1 feasibility (reconnaissance-level assessment, requiring <500 m grid "
        "resolution) and Stage 2 design (detailed site assessment, requiring <50 m grid "
        "resolution) with corresponding data requirements and methodologies. This dataset "
        "supports Stage 1 reconnaissance-level site screening."
    ),
    "CF Conventions": (
        "Climate and Forecast Conventions. A set of metadata conventions for earth science "
        "data that promote interoperability by standardizing variable names, units, and "
        "coordinate systems. This dataset follows CF conventions for NetCDF/HDF5 files."
    ),
    "PNNL": (
        "Pacific Northwest National Laboratory. A U.S. Department of Energy national "
        "laboratory that developed the original FVCOM model configurations and ran the "
        "hydrodynamic simulations used in this dataset."
    ),
    "WPTO": (
        "Water Power Technologies Office. The office within the U.S. Department of Energy's "
        "Office of Energy Efficiency and Renewable Energy that funds marine energy research "
        "and development, including this dataset."
    ),
    "VAP": (
        "Value-Added Product. A derived quantity calculated from the original model outputs "
        "to provide engineering-relevant metrics. Examples include mean current speed, power "
        "density, and water depth statistics."
    ),
}


def create_definitions_section(doc):
    """Create the Definitions section with technical terms, symbols, and acronyms."""
    add_heading(doc, "Definitions", level=1)

    intro_text = (
        "This section defines technical terms, mathematical symbols, and acronyms used "
        "throughout this document and in the dataset metadata."
    )
    add_body_paragraph(doc, intro_text)
    doc.add_paragraph()

    # Technical Terms subsection
    add_heading(doc, "Technical Terms", level=2)

    # Track table and figure numbers for this section
    table_counter = 1
    figure_counter = 1

    for term, definition in TECHNICAL_DEFINITIONS.items():
        # Add term heading
        p = doc.add_paragraph()
        p.add_run(f"{term}").bold = True

        # Handle both string (simple) and dict (complex with tables) formats
        if isinstance(definition, str):
            # Simple string definition
            p.add_run(f"\n{definition}")
        elif isinstance(definition, dict):
            # Complex definition with text and optional tables
            # Render initial text paragraphs
            text_content = definition.get("text", [])
            if isinstance(text_content, str):
                text_content = [text_content]

            for i, text_paragraph in enumerate(text_content):
                if i == 0:
                    # First paragraph continues after the term
                    p.add_run(f"\n{text_paragraph}")
                else:
                    # Subsequent paragraphs are new
                    add_body_paragraph(doc, text_paragraph)

            # Render tables if present
            tables = definition.get("tables", [])
            for table_spec in tables:
                doc.add_paragraph()  # Space before table
                render_table_from_spec(doc, table_spec, table_number=table_counter)
                table_counter += 1

            # Render text after tables if present
            text_after = definition.get("text_after_tables", [])
            if isinstance(text_after, str):
                text_after = [text_after]

            for text_paragraph in text_after:
                doc.add_paragraph()
                add_body_paragraph(doc, text_paragraph)

            # Render additional tables after text if present
            tables_after = definition.get("tables_after_text", [])
            for table_spec in tables_after:
                doc.add_paragraph()  # Space before table
                render_table_from_spec(doc, table_spec, table_number=table_counter)
                table_counter += 1

            # Render final text if present
            text_final = definition.get("text_final", [])
            if isinstance(text_final, str):
                text_final = [text_final]

            for text_paragraph in text_final:
                doc.add_paragraph()
                add_body_paragraph(doc, text_paragraph)

            # Render images if present
            images = definition.get("images", [])
            for image_spec in images:
                doc.add_paragraph()  # Space before image
                render_image_from_spec(doc, image_spec, figure_number=figure_counter)
                figure_counter += 1

        doc.add_paragraph()

    # Symbols subsection
    add_heading(doc, "Symbols", level=2)

    for symbol, definition in SYMBOL_DEFINITIONS.items():
        p = doc.add_paragraph()
        run = p.add_run(f"{symbol}")
        run.bold = True
        run.italic = True
        p.add_run(f"  —  {definition}")

    doc.add_paragraph()

    # Acronyms subsection
    add_heading(doc, "Acronyms", level=2)

    for acronym, definition in ACRONYM_DEFINITIONS.items():
        p = doc.add_paragraph()
        p.add_run(f"{acronym}").bold = True
        p.add_run(f"\n{definition}")
        doc.add_paragraph()

    doc.add_page_break()


def create_executive_summary(doc):
    """Create the executive summary section."""
    add_heading(doc, "Executive Summary", level=1)

    # Introduction paragraph
    intro = (
        "This document describes the variables from the WPTO High Resolution Tidal Hindcast "
        "dataset that are visualized on the Marine Energy Atlas. The dataset provides "
        "reconnaissance-level tidal energy resource characterization for five U.S. coastal "
        "regions, derived from high-resolution FVCOM hydrodynamic model simulations. These "
        "data support Stage 1 feasibility assessments per IEC 62600-201 standards."
    )
    add_body_paragraph(doc, intro)

    doc.add_paragraph()

    # Variables section
    add_heading(doc, "High Resolution Tidal Hindcast Variables", level=2)

    # Introduction paragraph
    variables_intro = "The following tidal energy quantities of interest are calculated from the original PNNL data and visualized on the NLR Marine Energy Atlas."
    add_body_paragraph(doc, variables_intro)

    # Create table with header row + data rows
    num_variables = len(ATLAS_VARIABLES)
    table = create_styled_table(doc, rows=num_variables + 1, cols=3, header_row=True)

    # Set column widths (Variable name narrower, Description wider)
    # table.columns[0].width = Inches(2.0)
    # table.columns[1].width = Inches(4.5)
    # table.columns[2].width = Inches(1.5)

    col_widths = [Inches(2.25), Inches(4.00), Inches(0.5)]

    # Populate header row
    header_cells = table.rows[0].cells
    header_cells[0].text = "Variable"
    header_cells[1].text = "Description"
    header_cells[2].text = "Units"

    for idx, width in enumerate(col_widths):
        header_cells[idx].width = width

    # Populate data rows
    for row_idx, (var_key, var_data) in enumerate(ATLAS_VARIABLES.items(), start=1):
        row = table.rows[row_idx]
        one_liner = var_data.get("one_liner", "")
        units = var_data["units"]

        # set_cell_text(row.cells[0], var_data["display_name"], bold=True)
        set_cell_text(row.cells[0], var_data["display_name"])
        set_cell_text(row.cells[1], one_liner)
        set_cell_text(row.cells[2], units)
        # # Variable name cell (bold)
        # var_cell = row.cells[0]
        # var_para = var_cell.paragraphs[0]
        # var_para.clear()
        # var_run = var_para.add_run(var_data["display_name"])
        # var_run.bold = True
        # var_run.font.name = "Arial"
        #
        # # Description cell
        # desc_cell = row.cells[1]
        # desc_cell.text = one_liner
        # desc_cell.font.name = "Arial"
        #
        # # Units cell
        # units_cell = row.cells[2]
        # units_cell.text = units
        # units_cell.font.name = "Arial"
        #
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_widths[idx]

    # Apply NREL styling with header and row striping
    style_table_with_striping(table, header_row=True)

    # Add table caption below the table with NLR styling
    add_table_caption(
        doc,
        "Summary of variables from the WPTO High Resolution Tidal Hindcast dataset available on the NLR Marine Energy Atlas",
        table_number=1,
    )

    doc.add_paragraph()

    # Location references section
    add_heading(doc, "Dataset Locations and Validation Studies", level=2)

    location_intro = (
        "The High Resolution Tidal Hindcast covers five U.S. coastal regions. Each location's "
        "model has been validated against field observations in peer-reviewed publications:"
    )
    add_body_paragraph(doc, location_intro)

    doc.add_paragraph()

    # Get bullet style for location references
    bullet_style = NREL_STYLES.get("bullet", "List Bullet")
    try:
        doc.styles[bullet_style]
    except KeyError:
        bullet_style = "List Bullet"

    # List locations with their full references and DOI hyperlinks
    for location, ref_keys in LOCATION_REFERENCES.items():
        # Location header
        p = doc.add_paragraph()
        p.add_run(f"{location}").bold = True

        # Add each reference with clickable DOI link parsed from the reference text
        for ref_key in ref_keys:
            try:
                full_ref = format_reference(ref_key)
                p = doc.add_paragraph(style=bullet_style)
                # Parse reference and make DOI a clickable link
                add_reference_with_doi_link(p, full_ref)
            except (ValueError, RuntimeError):
                pass

    doc.add_page_break()


def create_variable_complete_section(doc, var_key, var_data):
    """Create the complete documentation section for a variable (H2 under variable H1)."""
    # H2: Detailed Documentation
    add_heading(doc, "Detailed Documentation", level=2)

    # Units and internal name
    p = doc.add_paragraph()
    p.add_run("Units: ").bold = True
    p.add_run(var_data["units"])

    p = doc.add_paragraph()
    p.add_run("Internal Variable Name: ").bold = True
    run = p.add_run(var_data["column_name"])
    run.font.name = "Consolas"
    run.font.size = Pt(10)

    # Online documentation link
    doc_url = var_data.get("documentation_url", "")
    if doc_url:
        p = doc.add_paragraph()
        p.add_run("Online Documentation: ").bold = True
        add_hyperlink(p, doc_url, doc_url)

    # Complete description
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Description").bold = True

    # Split description into paragraphs, using body text style
    for para_text in var_data["complete_description"].split("\n\n"):
        if para_text.strip():
            add_body_paragraph(doc, para_text.strip())

    # Equation section
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Formula").bold = True

    # Add equation
    if var_key in EQUATION_BUILDERS:
        omml = EQUATION_BUILDERS[var_key]()
        add_omml_equation(doc, omml)

    # Add Where section
    add_where_section(doc, var_key, compact=False)

    # References for this variable
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("References").bold = True

    # Try NREL bullet style first, fall back to List Bullet
    bullet_style = NREL_STYLES.get("bullet", "List Bullet")
    try:
        doc.styles[bullet_style]
    except KeyError:
        bullet_style = "List Bullet"

    for ref_key in var_data["references"]:
        try:
            short_cite = format_citation(ref_key)
            doc.add_paragraph(short_cite, style=bullet_style)
        except (ValueError, RuntimeError):
            # Skip if reference not found in BibTeX
            pass

    doc.add_paragraph()  # Spacing


def create_document(use_nrel_template=True):
    """
    Create the complete Word document.

    Args:
        use_nrel_template: If True and NREL template exists, use it for styling.
                          Falls back to blank document if template not found.
    """
    # Try to use NREL template
    if use_nrel_template and NREL_TEMPLATE_PATH.exists():
        try:
            temp_docx_path = convert_dotx_to_docx(NREL_TEMPLATE_PATH)
            doc = Document(temp_docx_path)
            # Clear template content (keep styles AND section properties)
            # Preserve sectPr element - needed for page layout/table widths
            body = doc.element.body
            for element in body[:]:
                # Don't remove section properties - needed for page layout/table widths
                if element.tag != qn("w:sectPr"):
                    body.remove(element)
            print(f"Using NREL template: {NREL_TEMPLATE_PATH.name}")
        except Exception as e:
            print(f"Warning: Could not load NREL template ({e}), using default styles")
            doc = Document()
    else:
        doc = Document()
        # Set up default styles
        style = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(11)

    # Title - use NREL heading style or fall back to bold paragraph
    title_style = NREL_STYLES.get("heading1")
    title_text = "WPTO High Resolution Tidal Hindcast Dataset\nNLR Marine Energy Atlas Public Facing Variable Documentation"
    try:
        title = doc.add_paragraph(title_text, style=title_style)
    except KeyError:
        # Fall back to plain bold paragraph if Title style unavailable
        title = doc.add_paragraph()
        title.add_run(title_text).bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle with date
    # p = doc.add_paragraph()
    # p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # p.add_run("Marine Energy Atlas Variable Documentation").bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"Andrew Simms, Mike Lawson, Mary Serafin\nNational Laboratory of the Rockies\n{datetime.today().strftime('%B %-d, %Y')}"
    )

    doc.add_paragraph()

    # Executive Summary
    create_executive_summary(doc)

    # Overview (for website extraction)
    add_heading(doc, "Overview", level=1)
    intro_text = (
        "This document provides documentation for the user-facing variables displayed on the "
        "Marine Energy Atlas for the WPTO High Resolution Tidal Hindcast dataset. Each variable "
        "includes a short summary suitable for information popups, as well as complete technical "
        "documentation including formulas, calculation methodology, and references."
    )
    add_body_paragraph(doc, intro_text)

    # Data limitations
    limitations_text = (
        "Important limitations: This dataset is derived from numerical model simulations, not "
        "direct measurements. Results are based on a single hindcast year, which may not capture "
        "interannual variability in tidal energy resources. Model validation has been performed "
        "against observations at available measurement stations, but uncertainties exist, "
        "particularly in areas with complex bathymetry or limited observational data."
    )
    add_body_paragraph(doc, limitations_text)

    doc.add_page_break()

    # Definitions
    create_definitions_section(doc)

    # =========================================================================
    # Variable Documentation (each variable as H1 with Short and Detailed as H2)
    # =========================================================================
    for var_key, var_data in ATLAS_VARIABLES.items():
        # H1: Variable Name
        add_heading(doc, var_data["display_name"], level=1)

        # H2: Short Description (i) on Atlas
        create_variable_short_section(doc, var_key, var_data)

        # H2: Detailed Documentation
        create_variable_complete_section(doc, var_key, var_data)

        doc.add_page_break()

    # =========================================================================
    # SECTION 3: References (full bibliography using BibTeX/CSL)
    # =========================================================================
    add_heading(doc, "References", level=1)

    # Get appropriate reference style (NLR_Reference if available, else Normal)
    ref_style = get_reference_style(doc)

    for ref_key in BIBTEX_KEYS:
        try:
            # Get short form for label and full form for citation
            short_cite = format_citation(ref_key)
            full_cite = format_reference(ref_key)
            p = doc.add_paragraph(style=ref_style)
            p.add_run(f"{short_cite} ").bold = True
            p.add_run(full_cite)
            doc.add_paragraph()
        except (ValueError, RuntimeError):
            # Skip if reference not found in BibTeX
            pass

    return doc


# =============================================================================
# Markdown Generation for MkDocs
# =============================================================================

# LaTeX equations for each variable (matching the OMML equations)
LATEX_EQUATIONS = {
    "mean_current_speed": r"\bar{\bar{U}} = \text{mean}\left(\left[\text{mean}(U_{1,t}, \ldots, U_{N_\sigma,t}) \text{ for } t=1,\ldots,T\right]\right)",
    "mean_power_density": r"\bar{\bar{P}} = \text{mean}\left(\left[\text{mean}(P_{1,t}, \ldots, P_{N_\sigma,t}) \text{ for } t=1,\ldots,T\right]\right)",
    "p95_current_speed": r"U_{95} = P_{95}\left(\left[\max(U_{1,t}, \ldots, U_{N_\sigma,t}) \text{ for } t=1,\ldots,T\right]\right)",
    "min_water_depth": r"d_{\min} = \min_t(h + \zeta_t)",
    "max_water_depth": r"d_{\max} = \max_t(h + \zeta_t)",
    "grid_resolution": r"R = \frac{1}{3}(d_1 + d_2 + d_3)",
}


def to_kebab_case(text: str) -> str:
    """Convert text to kebab-case for anchor links."""
    # Convert to lowercase, replace spaces and special chars with hyphens
    text = text.lower()
    text = re.sub(r"[^a-z0-9]+", "-", text)
    text = text.strip("-")
    return text


def generate_markdown() -> str:
    """
    Generate MkDocs-compatible markdown documentation.

    Returns the markdown content as a string.
    """
    lines = []

    # Title
    lines.append("# Tidal Hindcast Variable Documentation")
    lines.append("")
    lines.append("Technical documentation for user-facing variables displayed on the Marine Energy Atlas for the WPTO High Resolution Tidal Hindcast dataset.")
    lines.append("")

    # Overview section
    lines.append("## Overview")
    lines.append("")
    lines.append(
        "This document provides documentation for the user-facing variables displayed on the "
        "Marine Energy Atlas for the WPTO High Resolution Tidal Hindcast dataset. Each variable "
        "includes a short summary suitable for information popups, as well as complete technical "
        "documentation including formulas, calculation methodology, and references."
    )
    lines.append("")
    lines.append("!!! warning \"Data Limitations\"")
    lines.append(
        "    This dataset is derived from numerical model simulations, not "
        "direct measurements. Results are based on a single hindcast year, which may not capture "
        "interannual variability in tidal energy resources. Model validation has been performed "
        "against observations at available measurement stations, but uncertainties exist, "
        "particularly in areas with complex bathymetry or limited observational data."
    )
    lines.append("")

    # Quick reference table
    lines.append("## Variable Quick Reference")
    lines.append("")
    lines.append("| Variable | Internal Name | Units | Description |")
    lines.append("|----------|---------------|-------|-------------|")
    for var_key, var_data in ATLAS_VARIABLES.items():
        anchor = to_kebab_case(var_data["display_name"])
        name_link = f"[{var_data['display_name']}](#{anchor})"
        lines.append(f"| {name_link} | `{var_data['column_name']}` | {var_data['units']} | {var_data['one_liner']} |")
    lines.append("")

    # Definitions section
    lines.append("## Definitions")
    lines.append("")
    lines.append(
        "This section defines technical terms, mathematical symbols, and acronyms used "
        "throughout this document and in the dataset metadata."
    )
    lines.append("")

    # Technical Terms
    lines.append("### Technical Terms")
    lines.append("")
    for term, definition in TECHNICAL_DEFINITIONS.items():
        lines.append(f"**{term}**")
        lines.append("")
        if isinstance(definition, str):
            lines.append(definition)
        elif isinstance(definition, dict):
            # Handle complex definitions with text arrays
            if "text" in definition:
                for text_block in definition["text"]:
                    lines.append(text_block)
                    lines.append("")
            # Include table data if present (sigma layers)
            if "tables" in definition:
                for table_spec in definition["tables"]:
                    lines.append(f"**{table_spec['caption']}**")
                    lines.append("")
                    headers = table_spec["headers"]
                    lines.append("| " + " | ".join(headers) + " |")
                    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                    # Get data rows
                    data = table_spec["data"]
                    num_rows = len(list(data.values())[0])
                    for i in range(num_rows):
                        row = [str(data[h][i]) for h in headers]
                        lines.append("| " + " | ".join(row) + " |")
                    lines.append("")
            if "text_after_tables" in definition:
                for text_block in definition["text_after_tables"]:
                    lines.append(text_block)
                    lines.append("")
            if "tables_after_text" in definition:
                for table_spec in definition["tables_after_text"]:
                    lines.append(f"**{table_spec['caption']}**")
                    lines.append("")
                    headers = table_spec["headers"]
                    lines.append("| " + " | ".join(headers) + " |")
                    lines.append("| " + " | ".join(["---"] * len(headers)) + " |")
                    data = table_spec["data"]
                    num_rows = len(list(data.values())[0])
                    for i in range(num_rows):
                        row = [str(data[h][i]) for h in headers]
                        lines.append("| " + " | ".join(row) + " |")
                    lines.append("")
            if "text_final" in definition:
                for text_block in definition["text_final"]:
                    lines.append(text_block)
                    lines.append("")
        lines.append("")

    # Symbols
    lines.append("### Symbols")
    lines.append("")
    lines.append("| Symbol | Definition |")
    lines.append("|--------|------------|")
    for symbol, definition in SYMBOL_DEFINITIONS.items():
        # Escape pipe characters in definition
        def_escaped = definition.replace("|", "\\|") if isinstance(definition, str) else str(definition).replace("|", "\\|")
        lines.append(f"| {symbol} | {def_escaped} |")
    lines.append("")

    # Acronyms
    lines.append("### Acronyms")
    lines.append("")
    lines.append("| Acronym | Definition |")
    lines.append("|---------|------------|")
    for acronym, definition in ACRONYM_DEFINITIONS.items():
        # Truncate long definitions for table, escape pipes
        def_text = definition if isinstance(definition, str) else str(definition)
        def_escaped = def_text.replace("|", "\\|")
        # Keep first sentence for brevity in table
        first_sentence = def_escaped.split(". ")[0] + "."
        lines.append(f"| {acronym} | {first_sentence} |")
    lines.append("")

    # Variable Documentation
    lines.append("## Variable Documentation")
    lines.append("")

    for var_key, var_data in ATLAS_VARIABLES.items():
        anchor = to_kebab_case(var_data["display_name"])

        # H3: Variable Name with anchor
        lines.append(f"### {var_data['display_name']} {{#{anchor}}}")
        lines.append("")

        # Metadata box
        lines.append(f"**Internal Name:** `{var_data['column_name']}`  ")
        lines.append(f"**Units:** {var_data['units']}")
        lines.append("")

        # One-liner (short description)
        lines.append(f"> {var_data['one_liner']}")
        lines.append("")

        # Complete description
        lines.append("#### Description")
        lines.append("")
        # Split description into paragraphs
        desc_paragraphs = var_data["complete_description"].split("\n\n")
        for para in desc_paragraphs:
            lines.append(para.strip())
            lines.append("")

        # Equation
        if var_key in LATEX_EQUATIONS:
            lines.append("#### Equation")
            lines.append("")
            lines.append(f"$$")
            lines.append(LATEX_EQUATIONS[var_key])
            lines.append(f"$$")
            lines.append("")

        # Where section (variable definitions)
        if var_key in VARIABLE_DEFINITIONS:
            lines.append("**Where:**")
            lines.append("")
            TEXT_DEFINITIONS = {
                "velocity_magnitude": "$U_{i,t} = \\sqrt{u^2 + v^2}$ — velocity magnitude at sigma layer $i$ at time $t$ (m/s)",
                "u_component": "$u$ = eastward velocity component (m/s), positive toward true east",
                "v_component": "$v$ = northward velocity component (m/s), positive toward true north",
                "power_density": "$P_{i,t} = \\frac{1}{2}\\rho U_{i,t}^3$ — power density at sigma layer $i$ at time $t$ (W/m²)",
                "n_sigma": "$N_\\sigma = 10$ sigma layers (terrain-following vertical layers)",
                "rho": "$\\rho = 1025$ kg/m³ (nominal seawater density)",
                "T": "$T$ = 1 year of hindcast data",
                "h": "$h$ = bathymetry depth below NAVD88 (m)",
                "zeta": "$\\zeta_t$ = sea surface elevation above NAVD88 at time $t$ (m)",
                "d_edges": "$d_1, d_2, d_3$ = geodesic distances between triangle vertices (m)",
                "max_sigma": "$\\max_\\sigma$ = maximum value across all sigma layers at each timestep",
                "p95": "$P_{95}$ = 95th percentile operator over the time series",
            }
            for def_key, _ in VARIABLE_DEFINITIONS[var_key]:
                if def_key in TEXT_DEFINITIONS:
                    lines.append(f"- {TEXT_DEFINITIONS[def_key]}")
            lines.append("")

        # References
        if "references" in var_data and var_data["references"]:
            lines.append("#### References")
            lines.append("")
            for ref_key in var_data["references"]:
                try:
                    full_cite = format_reference(ref_key)
                    short_cite = format_citation(ref_key)
                    lines.append(f"- **{short_cite}** {full_cite}")
                except (ValueError, RuntimeError):
                    lines.append(f"- {ref_key}")
            lines.append("")

        lines.append("---")
        lines.append("")

    # Full References section
    lines.append("## References")
    lines.append("")
    for ref_key in BIBTEX_KEYS:
        try:
            short_cite = format_citation(ref_key)
            full_cite = format_reference(ref_key)
            lines.append(f"**{short_cite}** {full_cite}")
            lines.append("")
        except (ValueError, RuntimeError):
            pass

    # Location-specific references
    lines.append("### Location-Specific Validation Studies")
    lines.append("")
    for location, refs in LOCATION_REFERENCES.items():
        lines.append(f"**{location}**")
        lines.append("")
        for ref_key in refs:
            try:
                full_cite = format_reference(ref_key)
                lines.append(f"- {full_cite}")
            except (ValueError, RuntimeError):
                lines.append(f"- {ref_key}")
        lines.append("")

    return "\n".join(lines)


def main():
    """Generate documentation in Word or Markdown format."""
    parser = argparse.ArgumentParser(
        description="Generate Marine Energy Atlas variable documentation"
    )
    parser.add_argument(
        "--format",
        choices=["docx", "markdown"],
        default="docx",
        help="Output format: 'docx' for Word document (default), 'markdown' for MkDocs"
    )
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help="Output file path (optional, defaults to output/ directory)"
    )
    args = parser.parse_args()

    output_dir = Path(__file__).parent / "output"
    output_dir.mkdir(exist_ok=True)

    if args.format == "markdown":
        # Generate markdown
        if args.output:
            output_path = Path(args.output)
        else:
            output_path = output_dir / "tidal_hindcast_variables.md"

        print("Generating Marine Energy Atlas Variable Documentation (Markdown)...")
        print(f"Output: {output_path}")

        markdown_content = generate_markdown()
        output_path.write_text(markdown_content, encoding="utf-8")

        print(f"\nMarkdown saved to: {output_path}")
    else:
        # Generate Word document (original behavior)
        if args.output:
            output_path = Path(args.output)
        else:
            output_path = output_dir / "marine_energy_atlas_variable_documentation.docx"

        print("Generating Marine Energy Atlas Variable Documentation (Word)...")
        print(f"Output: {output_path}")

        doc = create_document()
        doc.save(output_path)

        print(f"\nDocument saved to: {output_path}")

    print("\nVariables documented:")
    for var_key, var_data in ATLAS_VARIABLES.items():
        print(f"  - {var_data['display_name']} ({var_data['units']})")

    return output_path


if __name__ == "__main__":
    main()
